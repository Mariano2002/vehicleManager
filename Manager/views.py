from .forms import *
from django.shortcuts import render, redirect, get_object_or_404, HttpResponseRedirect
from django.views.decorators.csrf import csrf_exempt
import json
from .models import *
from django.contrib.auth.forms import UserCreationForm
from .forms import CreateUserForm
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.decorators import user_passes_test
import csv
from django.core import serializers
from django.urls import reverse
import re
import pandas as pd
from io import StringIO
import io
from datetime import date
from datetime import datetime, timedelta
import xlsxwriter
import math
from docx import Document
from docx.shared import Pt
import os
from django.conf import settings
from django.http import HttpResponse, Http404


@login_required(login_url="/login")
def download(request, path):
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-excel")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
            return response
    raise Http404

def update_balance():
    def days_between(d1, d2):
        return abs((d2 - d1).days)

    rentals_obj = rental.objects.all()
    for rental_i in rentals_obj:
        # balance = 0
        # # balance = balance+rental_i.deposit
        # for account_i in account.objects.all().filter(type='Car Rent', rental_id=rental_i.id):
        #     balance = balance+account_i.value
        # day_price = rental_i.rent/7
        # print(day_price)
        # print(days_between(rental_i.return_date, rental_i.hire_date)*day_price)
        # balance = balance-days_between(rental_i.return_date, rental_i.hire_date)*day_price
        # print(balance)
        # rental_i.balance = balance
        # rental_i.save()

        balance = 0
        # balance = balance+rental_i.deposit
        for account_i in account.objects.all().filter(type='Car Rent', rental_id=rental_i.id):
            balance = balance + account_i.value
        # print(balance)
        # day_price = rental_i.rent / 7
        # more_days = days_between(date.today(), rental_i.hire_date) - math.floor(days_between(date.today(), rental_i.hire_date)/7)
        # weeks_to_pay = 0
        #
        #
        # more_to_pay = math.floor(days_between(date.today(), rental_i.hire_date) / 7)
        #
        # weeks_to_pay += more_to_pay
        # print(weeks_to_pay)
        # if days_between(rental_i.return_date, date.today()) >= 7:
        #     balance = balance - weeks_to_pay * rental_i.rent
        # else:
        #     balance = balance - (weeks_to_pay * rental_i.rent) - (days_between(rental_i.return_date, date.today()) * day_price)
        print(balance)
        hire_date = rental_i.hire_date
        end_date = rental_i.return_date

        start = hire_date

        data = []
        day_price = rental_i.rent / 7
        while True:

            end = start + timedelta(days=6)

            if start > date.today():
                break
            elif end < end_date:
                data.append([start, end])
            else:
                data.append([start, end_date])
                break

            start = start + timedelta(days=7)

        for i in data:
            print(i)
            balance -= (days_between(i[-1], i[0]) + 1) * day_price
        print(balance)

        rental_i.balance = balance
        rental_i.save()

def update_vehicle_status():

    vehicles_obj = vehicle.objects.all().filter(status="On Rent")
    for i in vehicles_obj:
        try:
            print(rental.objects.get(vehicle_id=i.id, status="Active"))
        except:
            i.status = "Available"
            i.save()
    vehicles_obj = vehicle.objects.all().filter(status="Available")
    for i in vehicles_obj:
        try:
            print(rental.objects.get(vehicle_id=i.id, status="Active"))
            i.status = "On Rent"
            i.save()
        except:
            pass

def check_dates():
    pco = []
    mot = []
    road = []
    logbook = []
    total = []
    vehicles_obj = vehicle.objects.all()


    for i in vehicles_obj:
        mot_date = datetime.strptime(str(i.mot_expiration_date), '%Y-%m-%d')
        pco_date = datetime.strptime(str(i.pco_expiration_date), '%Y-%m-%d')
        road_date = datetime.strptime(str(i.road_tax_expiration_date), '%Y-%m-%d')
        logbook_date = datetime.strptime(str(i.LOGBOOK_first_date_of_registration), '%Y-%m-%d')
        if datetime.today() + timedelta(days=44) > mot_date:
            mot.append(i.id)
            total.append(i.id)
        if datetime.today() + timedelta(days=44) > pco_date:
            pco.append(i.id)
            total.append(i.id)
        if datetime.today() + timedelta(days=44) > road_date:
            road.append(i.id)
            total.append(i.id)
        if datetime.today() + timedelta(days=44) > logbook_date + timedelta(days=3652) :
            logbook.append(i.id)
            total.append(i.id)
    return pco, mot, road, logbook, total

@login_required(login_url="/login")
def home(request):
    if request.user.is_superuser:
        return render(request, 'home.html', {'deletes':deletes.objects.all(), 'super': True})
    else:
        return render(request, 'home.html', {})

def user_is_not_logged_in(user):
    return not user.is_authenticated

def user_is_admin(user):
    return user.is_superuser

def login_page(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        # print(user)
        if user is not None:
            login(request, user)
            last = lastSigned.objects.get(id=1)
            last.user_email = request.user.email
            last.username = request.user.username
            last.save()
            return redirect(home)
        else:
            messages.error(request, 'Username or password in incorrect')
            return render(request, 'login.html', {})
    else:
        # print(request.user)
        if request.user.is_anonymous:
            last_one = lastSigned.objects.get(id=1)
            return render(request, 'login.html', {"last_one":last_one})
        return redirect(home)


@user_passes_test(user_is_admin, login_url="/login")
def signup(request):
    levels = request.POST.getlist("levels")
    if "level1" in levels:
        level1 = "Yes"
    else:
        level1 = "No"
    if "level2" in levels:
        level2 = "Yes"
    else:
        level2 = "No"
    if "level3" in levels:
        level3 = "Yes"
    else:
        level3 = "No"
    if "level4" in levels:
        level4 = "Yes"
    else:
        level4 = "No"
    print(levels)
    form = CreateUserForm()

    if request.method == "POST":
        form = CreateUserForm(request.POST)
        if form.is_valid():
            form.save()
            permissions.objects.create(user=request.POST.get("username"), level1=level1, level2=level2, level3=level3, level4=level4)
            return redirect(login_page)
    context = {"form":form}
    return render(request, 'registration.html', context)

@login_required(login_url="/login")
def logout_page(request):
    print("here")
    logout(request)
    return redirect(login_page)






@login_required(login_url="/login")
def vehicles(request):
    update_vehicle_status()
    check_dates()
    pco, mot, road, logbook, total = check_dates()
    if request.method == "POST":
        update_balance()
        if request.POST.get('action') == "plate":
            print(request.POST.get("plate"))
            plate = request.POST.get("plate")
            vehicles_obj = vehicle.objects.all()
            for i in vehicles_obj:
                if plate not in i.plate:
                    vehicles_obj = vehicles_obj.exclude(id=i.id)

        elif request.POST.get('action') == "date":
            print(request.POST.get("date_type"))
            print(request.POST.get("daterange"))
            date_range = request.POST.get("daterange")
            type = request.POST.get("date_type")
            vehicles_obj = vehicle.objects.all()
            for i in vehicles_obj:
                start = datetime.strptime(date_range.split(" - ")[0], '%m/%d/%Y')
                end = datetime.strptime(date_range.split(" - ")[-1], '%m/%d/%Y')
                if type == "MOT":
                    date_now = datetime.strptime(str(i.mot_expiration_date), '%Y-%m-%d')
                elif type == "PCO":
                    date_now = datetime.strptime(str(i.pco_expiration_date), '%Y-%m-%d')
                elif type == "ROAD TAX":
                    date_now = datetime.strptime(str(i.road_tax_expiration_date), '%Y-%m-%d')

                if start <= date_now <= end:
                    pass
                else:
                    vehicles_obj = vehicles_obj.exclude(id=i.id)
        return render(request, 'vehicles.html', {'vehicles':vehicles_obj, 'pco':pco, 'mot':mot, 'road':road, 'logbook':logbook, 'total':total})
    else:
        update_balance()
        vehicles_obj = vehicle.objects.all()

        return render(request, 'vehicles.html', {'vehicles':vehicles_obj, 'pco':pco, 'mot':mot, 'road':road, 'logbook':logbook, 'total':total})

@login_required(login_url="/login")
def view_vehicle(request, vehicle_id):
    update_vehicle_status()
    update_balance()
    if request.method == "POST":

        if permissions.objects.get(user=request.user.username).level1 == "Yes":
            form = vehicleDocsForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                return redirect('view_vehicle', vehicle_id=vehicle_id)
            else:
                messages.error(request, 'Please use correct data!')
                return redirect('view_vehicle', vehicle_id=vehicle_id)
        else:
            return HttpResponse("Access Denied")
    else:
        vehicle_obj = vehicle.objects.all().filter(id=vehicle_id)
        doc_obj = vehicleDocs.objects.all().filter(vehicle_id=vehicle_id)
        rental_obj = rental.objects.all().filter(vehicle_id=vehicle_id)
        form = vehicleDocsForm
        return render(request, 'view_vehicle.html', {'vehicles':vehicle_obj, 'documents':doc_obj, 'form':form, 'rentals':rental_obj})

@login_required(login_url="/login")
def edit_vehicle(request, vehicle_id):
    if permissions.objects.get(user=request.user.username).level1 == "Yes":
        update_vehicle_status()

        if(request.method) == "POST":
            vehicle_i = vehicle.objects.all().filter(id=vehicle_id)[0]
            form = vehicleForm(data=request.POST, instance=vehicle_i)


            if form.is_valid():
                # print(form.errors)
                form.save()
                return redirect('view_vehicle', vehicle_id=vehicle_id)
            else:
                messages.error(request, 'Please use correct data!')
                return redirect('edit_vehicle', vehicle_id=vehicle_id)

        else:
            vehicle_obj = vehicle.objects.all().filter(id=vehicle_id)[0]
            form = vehicleForm
            owners = owner_name.objects.all()
            return render(request, 'edit_vehicle.html', {'vehicle':vehicle_obj, 'form':form, 'owners':owners})

    else:
        return HttpResponse("Access Denied")

@login_required(login_url="/login")
def add_vehicle(request):
    if permissions.objects.get(user=request.user.username).level1 == "Yes":
        if(request.method) == "POST":
            form = vehicleForm(request.POST)
            print(request.POST.get('owner_name'))
            try:
                owner_name.objects.create(name=request.POST.get('owner_name'))
            except:
                pass
            if form.is_valid():
                # print(form.errors)
                form.save()
                return redirect(vehicles)
            else:
                print(form.errors)
                messages.error(request, 'Please use correct data!')
                return redirect(add_vehicle)

        else:
            form = vehicleForm
            owners = owner_name.objects.all()
            return render(request, 'add_vehicle.html', {'form':form, 'owners':owners})
    else:
        return HttpResponse("Access Denied")

@login_required(login_url="/login")
def delete_vehicle(request, vehicle_id):

    if request.user.is_superuser:
        vehicle.objects.all().filter(id=vehicle_id)[0].delete()
        try:
            deletes.objects.get(link="delete_vehicle/"+str(vehicle_id)).delete()
            return redirect(deletes_view)
        except:
            pass
    else:
        try:
            deletes.objects.get(link="delete_vehicle/"+str(vehicle_id))
        except:
            deletes.objects.create(user=request.user.username, description="Delete vehicle with plate: "+vehicle.objects.get(id=vehicle_id).plate, link="delete_vehicle/"+str(vehicle_id))

    return redirect(vehicles)

@login_required(login_url="/login")
def delete_vehicleDoc(request, vehicleDoc_id):
    vehicleDoc_i = vehicleDocs.objects.all().filter(id=vehicleDoc_id)[0]
    vehicle_id = vehicleDoc_i.vehicle_id
    if request.user.is_superuser:
        vehicleDoc_i.delete()
        try:
            deletes.objects.get(link="delete_vehicleDoc/"+str(vehicleDoc_id)).delete()
            return redirect(deletes_view)
        except:
            pass
    else:
        try:
            deletes.objects.get(link="delete_vehicleDoc/"+str(vehicleDoc_id))
        except:
            deletes.objects.create(user=request.user.username, description="Delete vehicle Document with name: "+vehicleDocs.objects.get(id=vehicleDoc_id).file_name, link="delete_vehicleDoc/"+str(vehicleDoc_id))
    return redirect('view_vehicle', vehicle_id=vehicle_id)






@login_required(login_url="/login")
def drivers(request):
    update_balance()
    drivers_obj = driver.objects.all()

    return render(request, 'drivers.html', {'drivers':drivers_obj})

@login_required(login_url="/login")
def view_driver(request, driver_id):
    update_balance()
    if request.method == "POST":
        if permissions.objects.get(user=request.user.username).level2 == "Yes":

            form = driverDocsForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                return redirect('view_driver', driver_id=driver_id)
            else:
                messages.error(request, 'Please use correct data!')
                return redirect('view_driver', driver_id=driver_id)
        else:
            return HttpResponse("Access Denied")
    else:
        driver_obj = driver.objects.all().filter(id=driver_id)
        doc_obj = driverDocs.objects.all().filter(driver_id=driver_id)
        rental_obj = rental.objects.all().filter(driver_id=driver_id)
        form = driverDocsForm
        return render(request, 'view_driver.html', {'drivers':driver_obj, 'documents':doc_obj, 'form':form, 'rentals':rental_obj})

@login_required(login_url="/login")
def edit_driver(request, driver_id):
    if permissions.objects.get(user=request.user.username).level2 == "Yes":

        if(request.method) == "POST":
            driver_i = driver.objects.all().filter(id=driver_id)[0]
            form = driverForm(data=request.POST, instance=driver_i)


            if form.is_valid():
                # print(form.errors)
                form.save()
                return redirect('view_driver', driver_id=driver_id)
            else:
                messages.error(request, 'Please use correct data!')
                return redirect('edit_driver', driver_id=driver_id)

        else:
            driver_obj = driver.objects.all().filter(id=driver_id)[0]
            form = driverForm
            return render(request, 'edit_driver.html', {'driver':driver_obj, 'form':form})

    else:
        return HttpResponse("Access Denied")

@login_required(login_url="/login")
def add_driver(request):
    if permissions.objects.get(user=request.user.username).level2 == "Yes":
        if(request.method) == "POST":
            form = driverForm(request.POST)
            print(request.POST.get('birth_date'))
            if form.is_valid():
                # print(form.errors)
                form.save()
                return redirect(drivers)
            else:
                messages.error(request, 'Please use correct data!')
                return redirect(add_driver)

        else:
            form = driverForm
            return render(request, 'add_driver.html', {'form':form})
    else:
        return HttpResponse("Access Denied")

@login_required(login_url="/login")
def delete_driver(request, driver_id):

    if request.user.is_superuser:
        driver.objects.all().filter(id=driver_id)[0].delete()
        try:
            deletes.objects.get(link="delete_driver/"+str(driver_id)).delete()
            return redirect(deletes_view)
        except:
            pass
    else:
        try:
            deletes.objects.get(link="delete_driver/"+str(driver_id))
        except:
            deletes.objects.create(user=request.user.username, description="Delete driver with name: "+driver.objects.get(id=driver_id).name, link="delete_driver/"+str(driver_id))

    return redirect(drivers)

@login_required(login_url="/login")
def delete_driverDoc(request, driverDoc_id):

    driverDoc_i = driverDocs.objects.all().filter(id=driverDoc_id)[0]
    driver_id = driverDoc_i.driver_id
    if request.user.is_superuser:
        driverDoc_i.delete()
        try:
            deletes.objects.get(link="delete_driverDoc/"+str(driverDoc_id)).delete()
            return redirect(deletes_view)
        except:
            pass
    else:
        try:
            deletes.objects.get(link="delete_driverDoc/"+str(driverDoc_id))
        except:
            deletes.objects.create(user=request.user.username, description="Delete driver Document with name: "+driverDocs.objects.get(id=driverDoc_id).file_name, link="delete_driverDoc/"+str(driverDoc_id))
    return redirect('view_driver', driver_id=driver_id)






@login_required(login_url="/login")
def rentals(request):
    update_vehicle_status()
    update_balance()
    rentals_obj = rental.objects.all()

    return render(request, 'rentals.html', {'rentals':rentals_obj})

@login_required(login_url="/login")
def view_rental(request, rental_id):
    update_vehicle_status()
    update_balance()
    if request.method == "POST":
        if permissions.objects.get(user=request.user.username).level3 == "Yes":
            if request.POST.get("id") != None:
                rental_i = rental.objects.get(id=request.POST.get("id"))

                def days_between(d1, d2):
                    return abs((d2 - d1).days)

                day_price = rental_i.rent / 7
                total_to_pay = (days_between(rental_i.return_date, rental_i.hire_date)+1) * day_price

                output = io.BytesIO()
                workbook = xlsxwriter.Workbook(output)
                worksheet = workbook.add_worksheet()

                cell_format = workbook.add_format()
                cell_format.set_bold()

                worksheet.write('A1', 'Driver:', cell_format)
                worksheet.write('B1', rental_i.driver_full_name)
                worksheet.write('D1', "Vehicle:", cell_format)
                worksheet.write('E1', rental_i.vehicle_plate)

                worksheet.write('A2', 'Hire Date:', cell_format)
                worksheet.write('B2', str(rental_i.hire_date.strftime('%d-%m-%Y')))
                worksheet.write('A3', 'Return Date:', cell_format)
                worksheet.write('B3', str(rental_i.return_date.strftime('%d-%m-%Y')))
                worksheet.write('D2', "Rent:", cell_format)
                worksheet.write('E2', rental_i.rent)
                worksheet.write('D3', "Total to pay:", cell_format)
                worksheet.write('E3', total_to_pay)

                border_bottom = workbook.add_format({'bottom': 1})
                border_bottom.set_bold()

                row = 4
                col = 0
                for nr_max, i in enumerate(['Cash', 'Card', 'Account Transfer', 'Notes', 'Date']):
                    if i != "":
                        worksheet.write(row, col, i, border_bottom)
                    else:
                        worksheet.write(row, col, i, cell_format)
                    worksheet.set_column(row, col, 18)
                    col += 1

                accounts_obj = account.objects.all().filter(rental_id=str(request.POST.get("id")))
                print(accounts_obj)
                max_row = 0
                row = 5
                cash_balance = 0
                card_balance = 0
                account_balance = 0
                expenses = 0
                paid = 0
                for i in accounts_obj:
                    list_to_write = []
                    if i.payment_method == 'Cash':
                        col = 0
                        for el in list_to_write+["£"+str(i.value), '', '', i.notes, str(i.date.strftime('%d-%m-%Y'))]:
                            worksheet.write(row, col, el)
                            col += 1
                        row += 1
                        cash_balance += i.value
                        if i.notes == "Discount" or i.notes == "Expense":
                            expenses += i.value
                        else:
                            paid += i.value
                    if i.payment_method == 'Card':
                        col = 0
                        for el in list_to_write+['', "£"+str(i.value), '', i.notes, str(i.date.strftime('%d-%m-%Y'))]:
                            worksheet.write(row, col, el)
                            col += 1
                        row += 1
                        card_balance += i.value
                        if i.notes == "Discount" or i.notes == "Expense":
                            expenses += i.value
                        else:
                            paid += i.value
                    if i.payment_method == 'Account Transfer':
                        col = 0
                        for el in list_to_write+['', '', "£"+str(i.value), i.notes, str(i.date.strftime('%d-%m-%Y'))]:
                            worksheet.write(row, col, el)
                            col += 1
                        row += 1
                        account_balance += i.value
                        if i.notes == "Discount" or i.notes == "Expense":
                            expenses += i.value
                        else:
                            paid += i.value
                    if row > max_row:
                        max_row = row
                balances = {'cash': cash_balance, 'card': card_balance, 'account': account_balance}

                list_total = ["£" + str(balances['cash']), "£" + str(balances['card']),
                                               "£" + str(balances['account']), "", ""]

                border_top = workbook.add_format({'top': 1})
                border_top_bold = workbook.add_format({'top': 1})
                border_top_bold.set_bold()

                if max_row < 6:
                    max_row = 6
                row = max_row
                col = 0
                for i in list_total:
                    worksheet.write(row, col, i, border_top)
                    col += 1




                row += 1



                income = balances['cash'] + balances['card'] + balances['account']
                deposit = rental_i.deposit

                total = income-total_to_pay#+deposit
                worksheet.write(row+1, 0, "Total Income:", cell_format)
                worksheet.write(row+1, 1, "£"+str(income))
                worksheet.write(row+2, 0, "Total Paid:", cell_format)
                worksheet.write(row+2, 1, "£"+str(paid))
                worksheet.write(row+3, 0, "Total Expenses:", cell_format)
                worksheet.write(row+3, 1, "£"+str(expenses))
                worksheet.write(row+4, 0, "Deposit:", cell_format)
                worksheet.write(row+4, 1, "£"+str(deposit))
                worksheet.write(row+5, 0, "Balance:", cell_format)
                worksheet.write(row+5, 1, "£"+str(total))



                workbook.close()

                response = HttpResponse(content_type='application/vnd.ms-excel')

                response['Content-Disposition'] = f'attachment;filename="rental_{request.POST.get("id")}.xlsx"'

                response.write(output.getvalue())
                return response

            else:
                form = rentalDocsForm(request.POST, request.FILES)
                if form.is_valid():
                    form.save()
                    return redirect('view_rental', rental_id=rental_id)
                else:
                    messages.error(request, 'Please use correct data!')
                    return redirect('view_rental', rental_id=rental_id)
        else:
            return HttpResponse("Access Denied")
    else:

        update_balance()
        rental_obj = rental.objects.all().filter(id=rental_id)

        doc_obj = rentalDocs.objects.all().filter(rental_id=rental_id)

        accounts_rental = account.objects.all().filter(rental_id=rental_id)

        form = rentalDocsForm
        return render(request, 'view_rental.html', {'rentals':rental_obj, 'documents':doc_obj, 'accounts':accounts_rental, 'form':form})

@login_required(login_url="/login")
def add_rental(request):
    if permissions.objects.get(user=request.user.username).level3 == "Yes":
        update_vehicle_status()
        if(request.method) == "POST":

            updated_data = request.POST.copy()
            print(updated_data)
            print(updated_data['driver_id'])
            updated_data.update({'driver_full_name': driver.objects.all().filter(id=updated_data['driver_id'])[0].name+" "+driver.objects.all().filter(id=updated_data['driver_id'])[0].last_name})
            updated_data.update({'vehicle_plate': vehicle.objects.all().filter(id=updated_data['vehicle_id'])[0].plate})
            updated_data.update({'balance': 0})
            print(updated_data)
            form = rentalForm(data=updated_data)


            if updated_data['status'] == 'Active':
                vehicle_i = vehicle.objects.get(id=updated_data['vehicle_id'])
                vehicle_i.status = "On Rent"
                vehicle_i.save()
                if len(rental.objects.all().filter(vehicle_id=updated_data['vehicle_id'], status='Active')) != 0:
                    messages.error(request, 'This vehicle already has an active rental!')
                    return redirect(add_rental)

            if form.is_valid():
                # print(form.errors)
                form.save()
                return redirect(rentals)
            else:
                messages.error(request, 'Please use correct data!')
                return redirect(add_rental)

        else:
            form = rentalForm
            vehicles_obj = vehicle.objects.all().filter(status="Available")
            drivers_obj = driver.objects.all().filter(status="Active")
            return render(request, 'add_rental.html', {'form':form, 'vehicles':vehicles_obj, 'drivers':drivers_obj})

    else:
        return HttpResponse("Access Denied")
@login_required(login_url="/login")
def edit_rental(request, rental_id):
    if permissions.objects.get(user=request.user.username).level3 == "Yes":
        update_vehicle_status()

        if(request.method) == "POST":
            rental_i = rental.objects.all().filter(id=rental_id)[0]
            updated_data = request.POST.copy()
            print(updated_data)
            print(updated_data['driver_id'])
            updated_data.update({'driver_full_name': driver.objects.all().filter(id=updated_data['driver_id'])[0].name+" "+driver.objects.all().filter(id=updated_data['driver_id'])[0].last_name})
            updated_data.update({'vehicle_plate': vehicle.objects.all().filter(id=updated_data['vehicle_id'])[0].plate})
            updated_data.update({'balance': rental_i.balance})
            print(updated_data)
            form = rentalForm(data=updated_data, instance=rental_i)


            if updated_data['status'] == 'Active':
                vehicle_i = vehicle.objects.get(id=updated_data['vehicle_id'])
                vehicle_i.status = "On Rent"
                vehicle_i.save()
                if len(rental.objects.all().filter(vehicle_id=updated_data['vehicle_id'], status='Active')) != 0:
                    if rental.objects.all().filter(vehicle_id=updated_data['vehicle_id'], status='Active')[0].id != rental_id:
                        messages.error(request, 'This vehicle already has an active rental!')
                        return redirect(add_rental)
            else:
                vehicle_i = vehicle.objects.all().filter(id=updated_data['vehicle_id'], status='On Rent')
                if len(vehicle_i) != 0:
                    vehicle_i = vehicle_i[0]
                    vehicle_i.status = "Available"
                    vehicle_i.save()

            if form.is_valid():
                # print(form.errors)
                form.save()
                return redirect('view_rental', rental_id=rental_id)
            else:
                messages.error(request, 'Please use correct data!')
                return redirect('edit_rental', rental_id=rental_id)

        else:
            form = rentalForm
            vehicles_obj = vehicle.objects.all().filter(status="Available")
            vehicles_obj |= vehicle.objects.filter(id=rental.objects.get(id=rental_id).vehicle_id)
            drivers_obj = driver.objects.all().filter(status="Active")
            drivers_obj |= driver.objects.filter(id=rental.objects.get(id=rental_id).driver_id)
            rental_obj = rental.objects.all().filter(id=rental_id)[0]
            return render(request, 'edit_rental.html', {'rental':rental_obj, 'form':form, 'vehicles':vehicles_obj, 'drivers':drivers_obj})
    else:
        return HttpResponse("Access Denied")

@login_required(login_url="/login")
def delete_rental(request, rental_id):

    if request.user.is_superuser:
        rental.objects.all().filter(id=rental_id)[0].delete()
        try:
            deletes.objects.get(link="delete_rental/"+str(rental_id)).delete()
            return redirect(deletes_view)
        except:
            pass
    else:
        try:
            deletes.objects.get(link="delete_rental/"+str(rental_id))
        except:
            deletes.objects.create(user=request.user.username, description="Delete rental with ID : "+str(rental_id), link="delete_rental/"+str(rental_id))

    return redirect(rentals)


@login_required(login_url="/login")
def delete_rentalDoc(request, rentalDoc_id):

    rentalDoc_i = rentalDocs.objects.all().filter(id=rentalDoc_id)[0]
    rental_id = rentalDoc_i.rental_id
    if request.user.is_superuser:
        rentalDoc_i.delete()
        try:
            deletes.objects.get(link="delete_rentalDoc/"+str(rentalDoc_id)).delete()
            return redirect(deletes_view)
        except:
            pass
    else:
        try:
            deletes.objects.get(link="delete_rentalDoc/"+str(rentalDoc_id))
        except:
            deletes.objects.create(user=request.user.username, description="Delete rental Document with name: "+rentalDocs.objects.get(id=rentalDoc_id).file_name, link="delete_rentalDoc/"+str(rentalDoc_id))
    return redirect('view_rental', rental_id=rental_id)


@login_required(login_url="/login")
def print_rental_agreement(request, rental_id):
    if permissions.objects.get(user=request.user.username).level3 == "Yes":
        rental_i = rental.objects.get(id=rental_id)
        vehicle_i = vehicle.objects.get(id=rental_i.vehicle_id)
        driver_i = driver.objects.get(id=rental_i.driver_id)

        document = Document(r'Manager\static\Agreement Example.docx')
        dictionary = {
            "full_name": rental_i.driver_full_name,
            "address_driver": driver_i.address,
            "post_code": driver_i.postcode,
            "date_birth": str(driver_i.birth_date.strftime('%d-%m-%Y')),
            "tel_home": "",
            "license_number": driver_i.license_number,
            "expire_date": str(driver_i.LN_expiration_date.strftime('%d-%m-%Y')),
            "ni_number": driver_i.NI,
            "rent_week": rental_i.rent,
            "deposit_paid": rental_i.deposit,
            "tel_mob": driver_i.phone,
            "reg_mark": vehicle_i.plate,
            "make_veh": vehicle_i.make,
            "model_veh": vehicle_i.model,
            "color_veh": vehicle_i.color,
            "cylinder_cap": vehicle_i.cylinder_capacity,
            "fuel_type": vehicle_i.fuel_type,
            "mileage_out": vehicle_i.mileage,
            "date_out": rental_i.hire_date.strftime('%d-%m-%Y'),
            "time_o": str(datetime.now().strftime("%H:%M")),
            "date_in": rental_i.return_date.strftime('%d-%m-%Y'),
            "time_in": str(datetime.now().strftime("%H:%M")),
        }

        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for i in dictionary:
                            if i in paragraph.text:
                                paragraph.text = paragraph.text.replace(i, str(dictionary[i]))



        doc_io = io.BytesIO() # create a file-like object
        document.save(doc_io) # save data to file-like object
        doc_io.seek(0) # go to the beginning of the file-like object

        response = HttpResponse(doc_io.read())

        # Content-Disposition header makes a file downloadable
        response["Content-Disposition"] = "attachment; filename=Agreement.docx"

        # Set the appropriate Content-Type for docx file
        response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return response

    else:
        return HttpResponse("Access Denied")

@login_required(login_url="/login")
def print_rental_sheet(request, rental_id):
    if permissions.objects.get(user=request.user.username).level3 == "Yes":
        rental_i = rental.objects.get(id=rental_id)
        vehicle_i = vehicle.objects.get(id=rental_i.vehicle_id)
        driver_i = driver.objects.get(id=rental_i.driver_id)

        def days_between(d1, d2):
            return abs((d2 - d1).days)

        day_price = rental_i.rent / 7
        total_to_pay = days_between(rental_i.return_date, rental_i.hire_date) * day_price

        output = io.BytesIO()
        workbook = xlsxwriter.Workbook(output)
        worksheet = workbook.add_worksheet()

        cell_format_b_l = workbook.add_format({'left':1})
        cell_format_b_l.set_right(1)
        cell_format_b_l.set_top(1)
        cell_format_b_l.set_bottom(1)
        cell_format_b_l.set_bold()
        cell_format_b_l.set_bg_color('#d1c4e9')
        cell_format_b_h = workbook.add_format({'left':1})
        cell_format_b_h.set_right(1)
        cell_format_b_h.set_top(1)
        cell_format_b_h.set_bottom(1)
        cell_format_b_h.set_bold()
        cell_format_b_h.set_bg_color('#b39ddb')

        cell_format_l = workbook.add_format({'left':1})
        cell_format_l.set_right(1)
        cell_format_l.set_top(1)
        cell_format_l.set_bottom(1)
        cell_format_l.set_bg_color('#d1c4e9')
        cell_format_h = workbook.add_format({'left':1})
        cell_format_h.set_right(1)
        cell_format_h.set_top(1)
        cell_format_h.set_bottom(1)
        cell_format_h.set_bg_color('#b39ddb')

        cell_format_b = workbook.add_format({'left':2})
        cell_format_b.set_right(2)
        cell_format_b.set_top(2)
        cell_format_b.set_bottom(2)
        cell_format_b.set_bold()

        worksheet.write('A1', 'Ref:', cell_format_b_l)
        worksheet.write('B1', "", cell_format_l)
        worksheet.set_column(0, 0, 14)
        worksheet.write('A2', 'Car:', cell_format_b_h)
        worksheet.write('B2', vehicle_i.plate, cell_format_h)
        worksheet.set_column(1, 1, 11)
        worksheet.write('A3', 'Hire Start Date:', cell_format_b_l)
        worksheet.write('B3',  str(rental_i.hire_date.strftime('%d-%m-%Y')), cell_format_l)
        worksheet.write('A4', 'Deposit:', cell_format_b_h)
        if str(rental_i.deposit) == "0":
            worksheet.write('B4', "N/A", cell_format_h)
        else:
            worksheet.write('B4', rental_i.deposit, cell_format_h)

        worksheet.write('C1', '', cell_format_b_l)
        worksheet.write('D1', "", cell_format_l)
        worksheet.write('C2', '', cell_format_b_h)
        worksheet.write('D2', "", cell_format_h)
        worksheet.set_column(2, 2, 18)
        worksheet.write('C3', 'Start:', cell_format_b_l)
        worksheet.write('D3', str(rental_i.hire_date.strftime('%d-%m-%Y')), cell_format_l)
        worksheet.write('C4', 'End:', cell_format_b_h)
        worksheet.write('D4', str(rental_i.return_date.strftime('%d-%m-%Y')), cell_format_h)
        worksheet.set_column(3, 3, 11)
        worksheet.write('C5', 'Price:', cell_format_b_l)
        worksheet.write('D5', "£"+str(rental_i.rent)+" P/W", cell_format_l)

        worksheet.write('E1', 'Name:', cell_format_b_l)
        worksheet.write('F1', rental_i.driver_full_name, cell_format_l)
        worksheet.set_column(4, 4, 11)
        worksheet.write('E2', 'PostCode:', cell_format_b_h)
        worksheet.write('F2', driver_i.postcode, cell_format_h)
        worksheet.write('E3', 'Address:', cell_format_b_l)
        worksheet.write('F3', driver_i.address, cell_format_l)
        worksheet.write('E4', 'Phone:', cell_format_b_h)
        worksheet.write('F4', driver_i.phone, cell_format_h)
        worksheet.set_column(5, 5, 18)

        start = datetime.strptime(str(rental_i.hire_date), '%Y-%m-%d')
        end = datetime.strptime(str(rental_i.return_date), '%Y-%m-%d')
        rents = [{start.date():rental_i.rent}]
        next = start
        print(start)
        print(end)
        max = 2000
        while True:
            max -= 1
            if max == 0:
                break
            if next + timedelta(days=7) < end:
                next = next + timedelta(days=7)
                rents.append({next.date():rental_i.rent})
            else:
                rents[-1] = {next.date():rental_i.rent/7*(end-next).days}
                break
        print(rents)

        worksheet.write(10, 0, "Date", cell_format_b)
        worksheet.write(10, 1, "Amount", cell_format_b)
        worksheet.write(10, 2, "Description", cell_format_b)
        print(rents)
        if len(rents) > 1:
            rents1 = rents[:math.ceil(len(rents)/2)]
            rents2 = rents[math.ceil(len(rents)/2):]
            worksheet.write(10, 3, "Date", cell_format_b)
            worksheet.write(10, 4, "Amount", cell_format_b)
            worksheet.write(10, 5, "Description", cell_format_b)
        else:
            rents1 = rents
            rents2 = []

        row = 11
        for i in rents1:
            worksheet.write(row, 0, str((list(i.keys())[0]).strftime('%d %b')), cell_format_b)
            worksheet.write(row, 1, i[list(i.keys())[0]], cell_format_b)
            worksheet.write(row, 2, "", cell_format_b)
            row+=1
        row = 11
        for i in rents2:
            worksheet.write(row, 3, str((list(i.keys())[0]).strftime('%d %b')), cell_format_b)
            worksheet.write(row, 4, i[list(i.keys())[0]], cell_format_b)
            worksheet.write(row, 5, "", cell_format_b)
            row+=1


        merge_format = workbook.add_format()
        worksheet.merge_range(row+2, 0, row+5, 5, 'Notes:', merge_format)#

        workbook.close()

        response = HttpResponse(content_type='application/vnd.ms-excel')

        response['Content-Disposition'] = f'attachment;filename="rental_{rental_id}.xlsx"'

        response.write(output.getvalue())
        return response

    else:
        return HttpResponse("Access Denied")




@login_required(login_url="/login")
def add_account(request, type):
    if permissions.objects.get(user=request.user.username).level4 == "Yes":

        if request.method == "POST":
            post = request.POST.copy()  # to make it mutable
            print(request.POST.get("description"))
            if type == request.POST.get("type") == "Car Rent":
                rental_i = rental.objects.all().filter(id=request.POST.get("description"))[0]
                post["rental_id"] = request.POST.get("description")
                post["description"] = rental_i.vehicle_plate

            form = accountForm(data=post)
            print(post)
            if form.is_valid():
                # print(form.errors)
                form.save()

                return redirect(accounts)
            else:
                print(form.errors)
                messages.error(request, 'Please use correct data!')
                return redirect(add_account, type=type)
        else:
            form = accountForm

            today = date.today()
            print(today)
            if type == 'Expenditure':
                in_out = 'Outcome'
            else:
                in_out = 'Income'

            if type == "Car Rent":
                rental_obj = rental.objects.all().filter(status='Active')
                return render(request, 'add_account_2.html', {'form':form, 'type': type, 'today_date':today, 'in_out':in_out, 'rentals':rental_obj})
            else:
                return render(request, 'add_account_1.html', {'form':form, 'type': type, 'today_date':today, 'in_out':in_out})

    else:
        return HttpResponse("Access Denied")


@login_required(login_url="/login")
def accounts(request):

    if permissions.objects.get(user=request.user.username).level4 == "Yes":
        if request.method == "POST":
                type = request.POST.get('type')

                if type in ["Car Rent", 'Expenditure','PCO','Service & Repairs','Other Income','Car Sold']:
                    return redirect('add_account', type=type)
                elif type == None:
                    if request.user.is_superuser:
                        accounts_obj = account.objects.all()
                        vehicle_ids = request.POST.getlist('vehicles')
                        print(vehicle_ids)
                        # print(request.POST.getlist('drivers'))
                        driver_ids = request.POST.getlist('drivers')

                        # print(request.POST.getlist('owners'))
                        ownersi = request.POST.getlist('owners')

                        # print(request.POST.getlist('payment'))
                        payments = request.POST.getlist('payment')

                        types = request.POST.getlist('types')

                        date_range = request.POST.get('daterange')
                        vehicle_names = []
                        driver_names = []
                        for i in driver_ids:
                            driver_names.append(driver.objects.all().filter(id=i)[0].name+" "+driver.objects.all().filter(id=i)[0].last_name)
                        for i in vehicle_ids:
                            vehicle_names.append(vehicle.objects.all().filter(id=i)[0].plate)

                        if vehicle_names != [] or driver_names != [] or ownersi != []:
                            types = ["Car Rent"]

                        for i in accounts_obj:
                            if vehicle_ids != []:
                                try:
                                    if str(rental.objects.all().filter(id=i.rental_id)[0].vehicle_id) not in vehicle_ids:
                                        accounts_obj = accounts_obj.exclude(id=i.id)
                                except:
                                    accounts_obj = accounts_obj.exclude(id=i.id)


                            if driver_ids != []:
                                try:
                                    if str(rental.objects.all().filter(id=i.rental_id)[0].driver_id) not in driver_ids:
                                        accounts_obj = accounts_obj.exclude(id=i.id)
                                except:
                                    accounts_obj = accounts_obj.exclude(id=i.id)


                            if ownersi != []:
                                try:
                                    if vehicle.objects.all().filter(id=rental.objects.all().filter(id=i.rental_id)[0].vehicle_id)[0].owner_name not in ownersi:
                                        accounts_obj = accounts_obj.exclude(id=i.id)
                                except:
                                    accounts_obj = accounts_obj.exclude(id=i.id)

                            start = datetime.strptime(date_range.split(" - ")[0], '%m/%d/%Y')
                            end = datetime.strptime(date_range.split(" - ")[-1], '%m/%d/%Y')
                            date_now = datetime.strptime(str(i.date), '%Y-%m-%d')
                            if start <= date_now <= end:
                                pass
                            else:
                                accounts_obj = accounts_obj.exclude(id=i.id)


                            if payments != []:
                                if i.payment_method not in payments:
                                    accounts_obj = accounts_obj.exclude(id=i.id)

                            if types != []:
                                if i.type not in types:
                                    accounts_obj = accounts_obj.exclude(id=i.id)



                        if request.POST.get('action') == "filter":

                            owners = owner_name.objects.all()
                            vehicles = vehicle.objects.all()
                            drivers = driver.objects.all()
                            return render(request, 'accounts.html', {'accounts':accounts_obj, 'vehicles':vehicles, 'drivers':drivers, 'owners':owners})

                        elif request.POST.get('action') == "export":

                            output = io.BytesIO()
                            workbook = xlsxwriter.Workbook(output)
                            worksheet = workbook.add_worksheet()

                            cell_format = workbook.add_format()
                            cell_format.set_bold()

                            worksheet.write('A1', 'Income & Expenditure', cell_format)
                            worksheet.write('C1', 'Date:', cell_format)
                            worksheet.write('D1', "-".join([date_range.split(" - ")[0].split("/")[1], date_range.split(" - ")[0].split("/")[0] , date_range.split(" - ")[0].split("/")[2]]))
                            worksheet.write('E1', "-".join([date_range.split(" - ")[1].split("/")[1], date_range.split(" - ")[1].split("/")[0] , date_range.split(" - ")[1].split("/")[2]]))
                            if vehicle_names != [] or driver_names != [] or ownersi != []:
                                worksheet.write('G1', 'Vehicle(s):', cell_format)
                                if vehicle_names == []:
                                    worksheet.write('H1', "All")
                                else:
                                    worksheet.write('H1', ", ".join(vehicle_names))
                                worksheet.write('J1', 'Driver(s):', cell_format)
                                if driver_names == []:
                                    worksheet.write('K1', "All")
                                else:
                                    worksheet.write('K1', ", ".join(driver_names))
                                worksheet.write('M1', 'Owner(s):', cell_format)
                                if ownersi == []:
                                    worksheet.write('N1', "All")
                                else:
                                    worksheet.write('N1', ", ".join(ownersi))

                            list_types = []
                            if types == []:
                                types = ["Car Rent", 'Expenditure','PCO','Service & Repairs','Other Income','Car Sold']
                            for i in types:
                                list_types.append(i)
                                list_types.append("")
                                list_types.append("")
                                if i == "Expenditure":
                                    list_types.append("Outcome")
                                else:
                                    list_types.append("Income")
                                list_types.append("")
                                list_types.append("")
                                list_types.append("")


                            row = 2
                            col = 0
                            for i in list_types:
                                if col%7 == 0:
                                    worksheet.write(row, col, i, cell_format)
                                else:
                                    worksheet.write(row, col, i)
                                col +=1

                            headers = ['Description', 'Cash', 'Card', 'Account Transfer', 'Notes', 'Date', ""]

                            headers_full = []
                            for i in types:
                                headers_full = headers_full + headers



                            border_bottom = workbook.add_format({'bottom':1})
                            border_bottom.set_bold()

                            row = 3
                            col = 0
                            for nr_max, i in enumerate(headers_full):
                                if i != "":
                                    worksheet.write(row, col, i, border_bottom)
                                else:
                                    worksheet.write(row, col, i, cell_format)
                                worksheet.set_column(row, col, 18)
                                col +=1





                            balances = {}
                            max_row = 0
                            for nr, type in enumerate(types):
                                row = 4
                                cash_balance = 0
                                card_balance = 0
                                account_balance = 0
                                for i in accounts_obj:
                                    if i.type == type:
                                        list_to_write = []
                                        for u in range(0,nr*7):
                                            list_to_write.append("")
                                        if i.payment_method == 'Cash':
                                            col = 0
                                            for el in list_to_write+[i.description, "£"+str(i.value), '', '', i.notes, str(i.date.strftime('%d-%m-%Y'))]:
                                                worksheet.write(row, col, el)
                                                col += 1
                                            row += 1
                                            cash_balance += i.value
                                        if i.payment_method == 'Card':
                                            col = 0
                                            for el in list_to_write+[i.description, '', "£"+str(i.value), '', i.notes, str(i.date.strftime('%d-%m-%Y'))]:
                                                worksheet.write(row, col, el)
                                                col += 1
                                            row += 1
                                            card_balance += i.value
                                        if i.payment_method == 'Account Transfer':
                                            col = 0
                                            for el in list_to_write+[i.description, '', '', "£"+str(i.value), i.notes, str(i.date.strftime('%d-%m-%Y'))]:
                                                worksheet.write(row, col, el)
                                                col += 1
                                            row += 1
                                            account_balance += i.value
                                        if row > max_row:
                                            max_row = row

                                balances[type] = {'cash': cash_balance, 'card': card_balance, 'account': account_balance}



                            list_total = []
                            for i in balances:
                                list_total = list_total + ["Total:", "£"+str(balances[i]['cash']), "£"+str(balances[i]['card']), "£"+str(balances[i]['account']), "", "", '']





                            border_top = workbook.add_format({'top':1})
                            border_top_bold = workbook.add_format({'top':1})
                            border_top_bold.set_bold()

                            if max_row < 5:
                                max_row = 5
                            row = max_row
                            col = 0
                            for i in list_total:
                                if (col+1)%7 == 0 and col!=0:
                                    worksheet.write(row, col, i)
                                elif i == "Total:":
                                    worksheet.write(row, col, i, border_top_bold)
                                else:
                                    worksheet.write(row, col, i, border_top)
                                col +=1



                            row += 1


                            income = 0
                            outcome = 0
                            for i in balances:
                                if i == "Expenditure":
                                    outcome = outcome + balances[i]['cash'] + balances[i]['card'] + balances[i]['account']
                                else:
                                    income = income + balances[i]['cash'] + balances[i]['card'] + balances[i]['account']

                            for i in accounts_obj:
                                if i.type == "Car Rent":
                                    if i.notes == "Discount" or i.notes == "Expense":
                                        outcome = outcome + i.value
                                        income = income - i.value

                            total = income-outcome
                            worksheet.write(row+1, 0, "Total income:", cell_format)
                            worksheet.write(row+1, 1, "£"+str(income))
                            worksheet.write(row+2, 0, "Total expenses:", cell_format)
                            worksheet.write(row+2, 1, "£"+str(outcome))
                            worksheet.write(row+3, 0, "Profit & Loss:", cell_format)
                            worksheet.write(row+3, 1, "£"+str(total))

                            border_left_right = workbook.add_format({'left':1})
                            border_left_right.set_right(1)
                            border_left = workbook.add_format({'left':1})

                            row = 4
                            col = nr_max
                            for i in range(row, max_row):
                                for u in range(0, col):
                                    print(u)
                                    if (u+1)%7 == 0:
                                        worksheet.write(i, u, "", border_left_right)
                                    elif u == col-1:
                                        worksheet.write(i, u+1, "", border_left)




                            workbook.close()

                            response = HttpResponse(content_type='application/vnd.ms-excel')

                            response['Content-Disposition'] = f'attachment;filename="export_{date_range}.xlsx"'

                            response.write(output.getvalue())
                            return response



                    else:
                        return HttpResponse("Access Denied")
                else:
                    return redirect(accounts)
        else:
            accounts_obj = account.objects.all()

            n_days_ago = datetime.today() - timedelta(days=0)
            for i in accounts_obj:
                start = datetime.strptime(str(n_days_ago.date()), '%Y-%m-%d')
                end = datetime.strptime(str(date.today()), '%Y-%m-%d')
                date_now = datetime.strptime(str(i.date), '%Y-%m-%d')
                if start <= date_now <= end:
                    pass
                else:
                    accounts_obj = accounts_obj.exclude(id=i.id)
            owners = owner_name.objects.all()
            vehicles = vehicle.objects.all()
            drivers = driver.objects.all()
            return render(request, 'accounts.html', {'accounts':accounts_obj, 'vehicles':vehicles, 'drivers':drivers, 'owners':owners})

    else:
        return HttpResponse("Access Denied")


@login_required(login_url="/login")
def delete_account(request, account_id):

    if request.user.is_superuser:
        account.objects.all().filter(id=account_id)[0].delete()
        try:
            deletes.objects.get(link="delete_account/"+str(account_id)).delete()
            return redirect(deletes_view)
        except:
            pass
    else:
        try:
            deletes.objects.get(link="delete_account/"+str(account_id))
        except:
            deletes.objects.create(user=request.user.username, description="Delete transaction with ID : "+str(account_id), link="delete_account/"+str(account_id))

    return redirect(accounts)


@login_required(login_url="/login")
def delete_deleted(request, delete_id):

    if request.user.is_superuser:
        deletes.objects.all().filter(id=delete_id)[0].delete()

    return redirect(deletes_view)

@login_required(login_url="/login")
def deletes_view(request):
    return render(request, 'deletes.html', {'deletes': deletes.objects.all()})






# @login_required(login_url="/login")
# def owners(request):
#     owners_obj = owner.objects.all()
#
#     return render(request, 'owners.html', {'owners':owners_obj})
#
# @login_required(login_url="/login")
# def view_owner(request, owner_id):
#     update_vehicle_status()
#     update_balance()
#     if request.method == "POST":
#         if permissions.objects.get(user=request.user.username).level3 == "Yes":
#             if request.POST.get("id") != None:
#                 owner_i = owner.objects.get(id=request.POST.get("id"))
#
#                 def days_between(d1, d2):
#                     return abs((d2 - d1).days)
#
#                 day_price = owner_i.rent / 7
#                 total_to_pay = days_between(owner_i.return_date, owner_i.hire_date) * day_price
#
#                 output = io.BytesIO()
#                 workbook = xlsxwriter.Workbook(output)
#                 worksheet = workbook.add_worksheet()
#
#                 cell_format = workbook.add_format()
#                 cell_format.set_bold()
#
#                 worksheet.write('A1', 'Driver:', cell_format)
#                 worksheet.write('B1', owner_i.driver_full_name)
#                 worksheet.write('D1', "Vehicle:", cell_format)
#                 worksheet.write('E1', owner_i.vehicle_plate)
#
#                 worksheet.write('A2', 'Hire Date:', cell_format)
#                 worksheet.write('B2', str(owner_i.hire_date.strftime('%d-%m-%Y')))
#                 worksheet.write('A3', 'Return Date:', cell_format)
#                 worksheet.write('B3', str(owner_i.return_date.strftime('%d-%m-%Y')))
#                 worksheet.write('D2', "Rent:", cell_format)
#                 worksheet.write('E2', owner_i.rent)
#                 worksheet.write('D3', "Total to pay:", cell_format)
#                 worksheet.write('E3', total_to_pay)
#
#                 border_bottom = workbook.add_format({'bottom': 1})
#                 border_bottom.set_bold()
#
#                 row = 4
#                 col = 0
#                 for nr_max, i in enumerate(['Cash', 'Card', 'Account Transfer', 'Date']):
#                     if i != "":
#                         worksheet.write(row, col, i, border_bottom)
#                     else:
#                         worksheet.write(row, col, i, cell_format)
#                     worksheet.set_column(row, col, 18)
#                     col += 1
#
#                 accounts_obj = account.objects.all().filter(owner_id=str(request.POST.get("id")))
#                 print(accounts_obj)
#                 max_row = 0
#                 row = 5
#                 cash_balance = 0
#                 card_balance = 0
#                 account_balance = 0
#                 for i in accounts_obj:
#                     list_to_write = []
#                     if i.payment_method == 'Cash':
#                         col = 0
#                         for el in list_to_write+["£"+str(i.value), '', '', str(i.date.strftime('%d-%m-%Y'))]:
#                             worksheet.write(row, col, el)
#                             col += 1
#                         row += 1
#                         cash_balance += i.value
#                     if i.payment_method == 'Card':
#                         col = 0
#                         for el in list_to_write+['', "£"+str(i.value), '', str(i.date.strftime('%d-%m-%Y'))]:
#                             worksheet.write(row, col, el)
#                             col += 1
#                         row += 1
#                         card_balance += i.value
#                     if i.payment_method == 'Account Transfer':
#                         col = 0
#                         for el in list_to_write+['', '', "£"+str(i.value), str(i.date.strftime('%d-%m-%Y'))]:
#                             worksheet.write(row, col, el)
#                             col += 1
#                         row += 1
#                         account_balance += i.value
#                     if row > max_row:
#                         max_row = row
#                 balances = {'cash': cash_balance, 'card': card_balance, 'account': account_balance}
#
#                 list_total = ["£" + str(balances['cash']), "£" + str(balances['card']),
#                                                "£" + str(balances['account']), ""]
#
#                 border_top = workbook.add_format({'top': 1})
#                 border_top_bold = workbook.add_format({'top': 1})
#                 border_top_bold.set_bold()
#
#                 if max_row < 6:
#                     max_row = 6
#                 row = max_row
#                 col = 0
#                 for i in list_total:
#                     worksheet.write(row, col, i, border_top)
#                     col += 1
#
#
#
#
#                 row += 1
#
#
#
#                 outcome = total_to_pay
#                 income = balances['cash'] + balances['card'] + balances['account']
#                 deposit = owner_i.deposit
#
#                 total = income-outcome#+deposit
#                 worksheet.write(row+1, 0, "Total Paid:", cell_format)
#                 worksheet.write(row+1, 1, "£"+str(income))
#                 worksheet.write(row+2, 0, "Total to pay:", cell_format)
#                 worksheet.write(row+2, 1, "£"+str(outcome))
#                 worksheet.write(row+3, 0, "Deposit:", cell_format)
#                 worksheet.write(row+3, 1, "£"+str(deposit))
#                 worksheet.write(row+4, 0, "Balance:", cell_format)
#                 worksheet.write(row+4, 1, "£"+str(total))
#
#
#
#                 workbook.close()
#
#                 response = HttpResponse(content_type='application/vnd.ms-excel')
#
#                 response['Content-Disposition'] = f'attachment;filename="owner_{request.POST.get("id")}.xlsx"'
#
#                 response.write(output.getvalue())
#                 return response
#
#             else:
#                 form = ownerDocsForm(request.POST, request.FILES)
#                 if form.is_valid():
#                     form.save()
#                     return redirect('view_owner', owner_id=owner_id)
#                 else:
#                     messages.error(request, 'Please use correct data!')
#                     return redirect('view_owner', owner_id=owner_id)
#         else:
#             return HttpResponse("Access Denied")
#     else:
#
#         update_balance()
#         owner_obj = owner.objects.all().filter(id=owner_id)
#
#         doc_obj = ownerDocs.objects.all().filter(owner_id=owner_id)
#
#         accounts_owner = account.objects.all().filter(owner_id=owner_id)
#
#         form = ownerDocsForm
#         return render(request, 'view_owner.html', {'owners':owner_obj, 'documents':doc_obj, 'accounts':accounts_owner, 'form':form})
#
# @login_required(login_url="/login")
# def add_owner(request):
#     if permissions.objects.get(user=request.user.username).level3 == "Yes":
#         update_vehicle_status()
#         if(request.method) == "POST":
#
#             updated_data = request.POST.copy()
#             print(updated_data)
#             print(updated_data['driver_id'])
#             updated_data.update({'driver_full_name': driver.objects.all().filter(id=updated_data['driver_id'])[0].name+" "+driver.objects.all().filter(id=updated_data['driver_id'])[0].last_name})
#             updated_data.update({'vehicle_plate': vehicle.objects.all().filter(id=updated_data['vehicle_id'])[0].plate})
#             updated_data.update({'balance': 0})
#             print(updated_data)
#             form = ownerForm(data=updated_data)
#
#
#             if updated_data['status'] == 'Active':
#                 vehicle_i = vehicle.objects.get(id=updated_data['vehicle_id'])
#                 vehicle_i.status = "On Rent"
#                 vehicle_i.save()
#                 if len(owner.objects.all().filter(vehicle_id=updated_data['vehicle_id'], status='Active')) != 0:
#                     messages.error(request, 'This vehicle already has an active owner!')
#                     return redirect(add_owner)
#
#             if form.is_valid():
#                 # print(form.errors)
#                 form.save()
#                 return redirect(owners)
#             else:
#                 messages.error(request, 'Please use correct data!')
#                 return redirect(add_owner)
#
#         else:
#             form = ownerForm
#             vehicles_obj = vehicle.objects.all().filter(status="Available")
#             drivers_obj = driver.objects.all().filter(status="Active")
#             return render(request, 'add_owner.html', {'form':form, 'vehicles':vehicles_obj, 'drivers':drivers_obj})
#
#     else:
#         return HttpResponse("Access Denied")
# @login_required(login_url="/login")
# def edit_owner(request, owner_id):
#     if permissions.objects.get(user=request.user.username).level3 == "Yes":
#         update_vehicle_status()
#
#         if(request.method) == "POST":
#             owner_i = owner.objects.all().filter(id=owner_id)[0]
#             updated_data = request.POST.copy()
#             print(updated_data)
#             print(updated_data['driver_id'])
#             updated_data.update({'driver_full_name': driver.objects.all().filter(id=updated_data['driver_id'])[0].name+" "+driver.objects.all().filter(id=updated_data['driver_id'])[0].last_name})
#             updated_data.update({'vehicle_plate': vehicle.objects.all().filter(id=updated_data['vehicle_id'])[0].plate})
#             updated_data.update({'balance': owner_i.balance})
#             print(updated_data)
#             form = ownerForm(data=updated_data, instance=owner_i)
#
#
#             if updated_data['status'] == 'Active':
#                 vehicle_i = vehicle.objects.get(id=updated_data['vehicle_id'])
#                 vehicle_i.status = "On Rent"
#                 vehicle_i.save()
#                 if len(owner.objects.all().filter(vehicle_id=updated_data['vehicle_id'], status='Active')) != 0:
#                     if owner.objects.all().filter(vehicle_id=updated_data['vehicle_id'], status='Active')[0].id != owner_id:
#                         messages.error(request, 'This vehicle already has an active owner!')
#                         return redirect(add_owner)
#             else:
#                 vehicle_i = vehicle.objects.all().filter(id=updated_data['vehicle_id'], status='On Rent')
#                 if len(vehicle_i) != 0:
#                     vehicle_i = vehicle_i[0]
#                     vehicle_i.status = "Available"
#                     vehicle_i.save()
#
#             if form.is_valid():
#                 # print(form.errors)
#                 form.save()
#                 return redirect('view_owner', owner_id=owner_id)
#             else:
#                 messages.error(request, 'Please use correct data!')
#                 return redirect('edit_owner', owner_id=owner_id)
#
#         else:
#             form = ownerForm
#             vehicles_obj = vehicle.objects.all().filter(status="Available")
#             vehicles_obj |= vehicle.objects.filter(id=owner.objects.get(id=owner_id).vehicle_id)
#             drivers_obj = driver.objects.all().filter(status="Active")
#             drivers_obj |= driver.objects.filter(id=owner.objects.get(id=owner_id).driver_id)
#             owner_obj = owner.objects.all().filter(id=owner_id)[0]
#             return render(request, 'edit_owner.html', {'owner':owner_obj, 'form':form, 'vehicles':vehicles_obj, 'drivers':drivers_obj})
#     else:
#         return HttpResponse("Access Denied")
#
# @login_required(login_url="/login")
# def delete_owner(request, owner_id):
#
#     if request.user.is_superuser:
#         owner.objects.all().filter(id=owner_id)[0].delete()
#         try:
#             deletes.objects.get(link="delete_owner/"+str(owner_id)).delete()
#             return redirect(deletes_view)
#         except:
#             pass
#     else:
#         try:
#             deletes.objects.get(link="delete_owner/"+str(owner_id))
#         except:
#             deletes.objects.create(user=request.user.username, description="Delete owner with ID : "+str(owner_id), link="delete_owner/"+str(owner_id))
#
#     return redirect(owners)
#
#
# @login_required(login_url="/login")
# def delete_ownerDoc(request, ownerDoc_id):
#
#     ownerDoc_i = ownerDocs.objects.all().filter(id=ownerDoc_id)[0]
#     owner_id = ownerDoc_i.owner_id
#     if request.user.is_superuser:
#         ownerDoc_i.delete()
#         try:
#             deletes.objects.get(link="delete_ownerDoc/"+str(ownerDoc_id)).delete()
#             return redirect(deletes_view)
#         except:
#             pass
#     else:
#         try:
#             deletes.objects.get(link="delete_ownerDoc/"+str(ownerDoc_id))
#         except:
#             deletes.objects.create(user=request.user.username, description="Delete owner Document with name: "+ownerDocs.objects.get(id=ownerDoc_id).file_name, link="delete_ownerDoc/"+str(ownerDoc_id))
#     return redirect('view_owner', owner_id=owner_id)
