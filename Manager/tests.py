from docx import Document
from docx.shared import Pt

document = Document(r'C:\Users\Mariano\Desktop\Agreement Example.docx')
dictionary = {
"full_name":"Mariano Baci",
"address_driver":"Jahja Ballhysa",
"post_code":"2001",
"date_birth":"22-01-2002",
"tel_home":"12345",
"license_number":"23534612345",
"expire_date":"23-03-2024",
"ni_number":"123423",
"rent_week":"70",
"deposit_paid":"50",
"tel_mob":"12345678",
"reg_mark":"SER 2FW",
"make_veh":"Ford",
"model_veh":"Focus",
"date_out":"24-01-2021",
"time_o":"11:35",
"date_in":"24-01-2021",
"time_in":"23:00",
}

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

for paragraph in document.paragraphs:
    for i in dictionary:
        if i in paragraph.text:
            paragraph.text = 'new text containing ocean'

for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for i in dictionary:
                    if i in paragraph.text:
                        paragraph.text = paragraph.text.replace(i, dictionary[i])
document.save(r'C:\Users\Mariano\Desktop\Agreement Example1.docx')